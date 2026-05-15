import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_header(doc, title):
    # Add title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r_title = p_title.add_run(f"{title}")
    r_title.bold = True
    r_title.font.size = Pt(16)
    
    # Add Project Details
    p_details = doc.add_paragraph()
    p_details.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_details.add_run("Project Title: Medical Clinic Database System\n").bold = True
    p_details.add_run("Group Members: Idrees Zazai, Farman Afridi\n").bold = True
    p_details.add_run("Instructor: Ali Hassan\n").bold = True
    
    doc.add_paragraph() # Spacing

def create_milestone_2():
    doc = Document()
    add_header(doc, "Milestone 2: ERD Design & Normalization")
    
    doc.add_heading('Step 1 & 2 — Apply Normalization & Remove Duplicates', level=1)
    
    doc.add_heading('1NF (First Normal Form)', level=2)
    doc.add_paragraph("Issue Found: The full_name fields in USERS, PATIENTS, and DOCTORS were not atomic. The prescribed_medicines field in VISITS contained multiple medicines.")
    doc.add_paragraph("Changes Made: Split full_name into first_name and last_name across all relevant tables. Removed prescribed_medicines from VISITS and created a new PRESCRIPTIONS table.")
    doc.add_paragraph("Why: To comply with 1NF, every column must hold a single, indivisible value.")
    
    doc.add_heading('2NF (Second Normal Form)', level=2)
    doc.add_paragraph("Issue Found: None.")
    doc.add_paragraph("Changes Made: None.")
    doc.add_paragraph("Why: Partial dependencies only exist with composite primary keys. All tables use single-column surrogate primary keys. Thus, being in 1NF guarantees 2NF.")
    
    doc.add_heading('3NF (Third Normal Form)', level=2)
    doc.add_paragraph("Issue Found: None.")
    doc.add_paragraph("Changes Made: None.")
    doc.add_paragraph("Why: All attributes depend strictly on the primary key, avoiding transitive dependencies. The schema is successfully in 3NF.")
    
    doc.add_heading('Step 3 — Updated ERD', level=1)
    doc.add_paragraph("The ER Diagram has been updated to reflect the new PRESCRIPTIONS table and atomic name fields. Please refer to the updated_erd.png image in the ERD directory.")
    
    doc.save('Milestones/Milestone 2.docx')

def create_milestone_3():
    doc = Document()
    add_header(doc, "Milestone 3: Dataset Preprocessing")
    
    doc.add_heading('Step 1 — Prepare Your Dataset', level=1)
    doc.add_paragraph("A Python script was utilized to generate a realistic and structured synthetic dataset. We generated exactly 60 rows of data for all 8 core tables (USERS, PATIENTS, DOCTORS, MEDICAL_SERVICES, VISITS, PRESCRIPTIONS, VISIT_SERVICES, BILLS) using randomized names, dates, and medical terminology.")
    
    doc.add_heading('Step 2 — Define Your Dataflow', level=1)
    doc.add_paragraph("1. Data Entry Point: The receptionist registers new patients in the PATIENTS table.")
    doc.add_paragraph("2. Core Operational Flow: The receptionist schedules a visit by connecting PATIENTS to DOCTORS within the VISITS table.")
    doc.add_paragraph("3. Dependent Streams: From VISITS, doctors issue PRESCRIPTIONS and order MEDICAL_SERVICES (tracked in VISIT_SERVICES).")
    doc.add_paragraph("4. Financial Processing: All services linked to a visit calculate a total_amount, stored in the BILLS table.")
    doc.add_paragraph("5. Outputs: The system generates Patient Invoices, Medical History Reports, and Operational Reports.")
    
    doc.add_heading('Step 3 — Export Clean CSV Files', level=1)
    doc.add_paragraph("The generated dataset was successfully exported as 8 clean CSV files into the Data directory.")
    
    doc.save('Milestones/Milestone 3.docx')

def create_milestone_4():
    doc = Document()
    add_header(doc, "Milestone 4: Database Setup (DDL)")
    
    doc.add_heading('DDL Scripts and Constraints', level=1)
    doc.add_paragraph("The schema.sql script was created mapping our normalized schema to MySQL. It includes:")
    doc.add_paragraph("- AUTO_INCREMENT PRIMARY KEY constraints on all tables.")
    doc.add_paragraph("- FOREIGN KEY constraints linking all relationships with ON DELETE CASCADE.")
    doc.add_paragraph("- NOT NULL applied to essential fields.")
    doc.add_paragraph("- UNIQUE constraint on username.")
    doc.add_paragraph("- CHECK constraints ensuring data validity (e.g., price >= 0, gender IN ('Male', 'Female', 'Other'), role IN ('Admin', 'Receptionist')).")
    
    doc.add_heading('Indexes for Performance', level=1)
    doc.add_paragraph("We implemented CREATE INDEX statements on all foreign key columns to ensure foreign key lookups are optimized. Additional indexes were placed on frequently queried columns such as patient last_name, doctor specialty, and visit_date.")
    
    doc.add_heading('EER Diagram Verification', level=1)
    doc.add_paragraph("The DDL script has been written such that MySQL Workbench's Reverse Engineer feature perfectly generates the final EER Diagram.")
    
    doc.save('Milestones/Milestone 4.docx')

def create_milestone_5():
    doc = Document()
    add_header(doc, "Milestone 5: Data Population (DML)")
    
    doc.add_heading('Data Population Strategy', level=1)
    doc.add_paragraph("Instead of LOAD DATA INFILE, a custom Python script converted our 8 CSV files into a comprehensive insert_data.sql script. This file contains over 400 lines of pure INSERT INTO statements to safely populate the database without triggering MySQL file permission restrictions.")
    
    doc.add_heading('Validation Queries', level=1)
    doc.add_paragraph("A validation_queries.sql script was developed to fulfill all final assignment requirements:")
    doc.add_paragraph("1. UPDATE Operation: Updated a bill's status to 'Paid' using a WHERE clause.")
    doc.add_paragraph("2. DELETE Operation: Deleted an old visit using a WHERE clause.")
    doc.add_paragraph("3. Row Count: A UNION ALL query executing COUNT(*) across all 8 tables.")
    doc.add_paragraph("4. NULL Checks: Queries targeting first_name, visit_date, and total_amount to ensure data completeness.")
    doc.add_paragraph("5. JOIN-Based FK Integrity: A LEFT JOIN query checking for orphaned records, guaranteeing foreign key constraints are fully intact.")
    
    doc.save('Milestones/Milestone 5.docx')

if __name__ == "__main__":
    os.makedirs('Milestones', exist_ok=True)
    create_milestone_2()
    create_milestone_3()
    create_milestone_4()
    create_milestone_5()
    print("All Milestone Word documents created successfully in the Milestones/ directory!")
